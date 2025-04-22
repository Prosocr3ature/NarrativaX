import os
import json
import requests
import zipfile
import random
import replicate
import time

from io import BytesIO
from tempfile import NamedTemporaryFile
from PIL import Image
from docx import Document
from gtts import gTTS, gTTSError

import streamlit as st
from streamlit_sortables import sort_items

# ====================
# CONFIG & CONSTANTS
# ====================
MAX_TOKENS = 1800
IMAGE_SIZE = (768, 1024)
SAFE_LOADING_MESSAGES = [
    "📚 Plotting character arcs...", "🎨 Mixing narrative tones...",
    "💡 Crafting plot twists...", "🌌 Building worlds...",
    "🔥 Igniting conflicts...", "💔 Developing relationships..."
]

GENRES = {
    "📖 Literary":    ["Thoughtful", "Reflective", "Historical Fiction", "Experimental", "Magical Realism"],
    "💖 Romance":     ["Sensual", "Emotional", "Paranormal Romance", "Romantic Comedy", "Drama"],
    "🔞 Adult":       ["Erotic", "Softcore", "Hardcore", "BDSM", "Fetish"],
    "🚀 Sci‑Fi":      ["Space Opera", "Cyberpunk", "Dystopian", "Time Travel", "Military Sci‑Fi"],
    "🪄 Fantasy":     ["Epic", "Urban Fantasy", "Dark Fantasy", "Fairy Tale", "Steampunk"],
    "🔪 Thriller":    ["Suspense", "Psychological", "Crime", "Espionage", "Noir"],
    "🕵️ Mystery":    ["Detective", "Whodunit", "Cozy Mystery", "Forensic", "Locked‑Room"],
    "👻 Horror":      ["Supernatural", "Slasher", "Paranormal", "Gore", "Haunted House"],
    "🏰 Historical":  ["Period Drama", "Alternate History", "Medieval", "Victorian", "World War"],
    "💼 Business":    ["Entrepreneurship", "Leadership", "Finance", "Strategy", "Startups"],
    "🌱 Self‑Help":   ["Motivational", "Mindfulness", "Productivity", "Wellness", "Habits"],
    "👶 Children":    ["Bedtime Story", "Young Adult", "Middle Grade", "Picture Book", "Coming‑of‑Age"]
}

TONES = {
    "😊 Wholesome":     "Uplifting, positive, family‑friendly",
    "😈 Explicit":      "Graphic, raw, adult‑oriented",
    "🤔 Philosophical": "Contemplative, existential, deep",
    "😄 Humorous":      "Witty, lighthearted, satirical",
    "😱 Dark":          "Gritty, intense, disturbing",
    "💡 Educational":   "Informative, structured, factual",
    "💔 Melancholic":   "Sad, poignant, bittersweet",
    "🔮 Mystical":      "Ethereal, mysterious, dreamlike",
    "🔥 Passionate":    "Fervent, intense, ardent",
    "🦄 Quirky":        "Eccentric, whimsical, offbeat"
}

IMAGE_MODELS = {
    "🎨 Realistic Vision":    "lucataco/realistic-vision-v5.1:…",
    "🔥 Reliberate NSFW":     "asiryan/reliberate-v3:…",
    "🍑 Uber‑Realistic Porn": "ductridev/uber-realistic-porn-merge-urpm-1:1cca487c3bfe167e987fc3639477cf2cf617747cd38772421241b04d27a113a8",
}

MODELS = {
    "🧠 MythoMax": "gryphe/mythomax-l2-13b",
    "🐬 Dolphin":  "cognitivecomputations/dolphin-mixtral",
    "🤖 OpenChat": "openchat/openchat-3.5-0106"
}

PRESETS = {
    "Vanilla": 0,
    "NSFW":    50,
    "Hardcore":100
}


# ====================
# STATE INIT
# ====================
def initialize_state():
    defaults = {
        "last_saved":           None,
        "book":                 {},
        "chapter_order":        [],
        "outline":              "",
        "cover":                None,
        "image_cache":          {},
        "characters":           [],  # list of dicts: {"desc": str, "img": PIL.Image or None}
        "gen_progress":         {},
        "selected_genre":       None,
        "selected_subgenres":   [],
        "selected_tone":        None,
        "selected_model":       list(MODELS.keys())[0],
        "selected_image_model": list(IMAGE_MODELS.keys())[0],
        "content_preset":       "Vanilla",
        "explicit_level":       PRESETS["Vanilla"],
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v


# ====================
# ENV CHECK
# ====================
def validate_environment():
    missing = [k for k in ("OPENROUTER_API_KEY", "REPLICATE_API_TOKEN") if k not in st.secrets]
    if missing:
        st.error(f"Missing API keys: {', '.join(missing)}")
        st.stop()
    try:
        replicate.Client(api_token=st.secrets["REPLICATE_API_TOKEN"])
    except Exception as e:
        st.error(f"Replicate auth failed: {e}")
        st.stop()


# ====================
# ZIP EXPORT
# ====================
def create_export_zip():
    buf = BytesIO()
    with zipfile.ZipFile(buf, "w") as z:
        doc = Document()
        for title in st.session_state.chapter_order:
            doc.add_heading(title, level=1)
            doc.add_paragraph(st.session_state.book[title])
        with NamedTemporaryFile(suffix=".docx") as tmp:
            doc.save(tmp.name)
            z.write(tmp.name, "book.docx")
        for sec, img in st.session_state.image_cache.items():
            with NamedTemporaryFile(suffix=".jpg") as tmp:
                img.save(tmp.name)
                z.write(tmp.name, f"images/{sec}.jpg")
    buf.seek(0)
    return buf


# ====================
# AI / IMG CALLS
# ====================
def call_openrouter(prompt, model_key):
    headers = {
        "Authorization": f"Bearer {st.secrets['OPENROUTER_API_KEY']}",
        "Content-Type": "application/json"
    }
    explicit = f"[Explicit Level: {st.session_state.explicit_level}/100] "
    payload = {
        "model":       MODELS[model_key],
        "messages":    [{"role": "user", "content": explicit + prompt}],
        "max_tokens":  MAX_TOKENS,
        "temperature": 0.7 + st.session_state.explicit_level * 0.002
    }
    try:
        r = requests.post(
            "https://openrouter.ai/api/v1/chat/completions",
            headers=headers, json=payload, timeout=30
        )
        r.raise_for_status()
        return r.json()["choices"][0]["message"]["content"]
    except Exception as e:
        st.error(f"API Error: {e}")
        return ""


def generate_image(prompt, model_key, section):
    version = IMAGE_MODELS[model_key]
    neg = "text, watermark" if "NSFW" in model_key else ""
    add = f", explicit {st.session_state.explicit_level}/100" if st.session_state.explicit_level > 0 else ""
    try:
        out = replicate.run(version, input={
            "prompt":          prompt + add,
            "width":           IMAGE_SIZE[0],
            "height":          IMAGE_SIZE[1],
            "negative_prompt": neg
        })
        url = out[0] if isinstance(out, list) else out
        resp = requests.get(url, timeout=15)
        resp.raise_for_status()
        img = Image.open(BytesIO(resp.content))
        st.session_state.image_cache[section] = img
        return img
    except Exception as e:
        st.error(f"Image generation failed: {e}")
        return None


# ====================
# BOOK & CHARACTER GENERATION
# ====================
def generate_book_content():
    # reset previous content
    st.session_state.book = {}
    st.session_state.chapter_order = []
    st.session_state.image_cache = {}
    st.session_state.cover = None
    st.session_state.characters = []

    with st.spinner("📖 Crafting Your Masterpiece..."):
        try:
            st.write(random.choice(SAFE_LOADING_MESSAGES))
            gp = st.session_state.gen_progress

            # Outline
            outline_prompt = (
                f"Create a {gp['chapters']}-chapter outline for "
                f"{gp['genre']} ({', '.join(gp['subgenres'])}) "
                f"with {gp['tone']} tone. Seed: {gp['prompt']}"
            )
            outline = call_openrouter(outline_prompt, gp["model"])
            st.session_state.outline = outline

            # Chapters
            lines = [l.strip() for l in outline.split("\n") if l.strip()]
            # assume first lines are header, drop until chapters list
            chapters_list = [l for l in lines if l.lower().startswith("chapter")]
            for chap in chapters_list:
                st.write(f"📝 Writing {chap}…")
                content = call_openrouter(f"Expand this chapter: {chap}", gp["model"])
                st.session_state.book[chap] = content
                st.session_state.chapter_order.append(chap)

            # Cover
            st.write("🎨 Painting the cover…")
            cover = generate_image(
                f"Cover for {gp['prompt']}, {gp['genre']} ({', '.join(gp['subgenres'])}), {gp['tone']}",
                gp["img_model"], "cover"
            )
            st.session_state.cover = cover

            # Automatic character generation
            st.write("🧑 Generating characters…")
            chars_text = call_openrouter(
                f"Generate 3 characters for story seed: {gp['prompt']} in {gp['genre']}, tone {gp['tone']}",
                gp["model"]
            )
            for c in [c.strip() for c in chars_text.split("\n\n") if c.strip()]:
                st.session_state.characters.append({"desc": c, "img": None})

        except Exception as e:
            st.error(f"Generation error: {e}")


# ====================
# MAIN UI
# ====================
def main_interface():
    st.set_page_config(page_title="NarrativaX", page_icon="📚", layout="wide")
    initialize_state()
    validate_environment()

    # CSS tweaks
    st.markdown("""
    <style>
      .block-container { padding: 8px 16px; }
      .logo-container { text-align: center; margin: 12px 0; }
      .stHeader, .stSubheader { margin: 4px 0 !important; }
      .stButton > button { margin: 4px !important; }
      .stExpander > .stMarkdown { margin-bottom: 4px !important; }
    </style>
    """, unsafe_allow_html=True)

    # Sidebar
    with st.sidebar:
        st.markdown("### 🔧 Session")
        if st.session_state.last_saved:
            st.caption(f"⏱️ Last saved: {time.strftime('%Y-%m-%d %H:%M', time.localtime(st.session_state.last_saved))}")
        c1, c2, c3 = st.columns(3)
        with c1:
            if st.button("💾 Save"):
                with open("session.json", "w") as f:
                    json.dump({
                        "book":     st.session_state.book,
                        "chapters": st.session_state.chapter_order,
                        "characters": st.session_state.characters
                    }, f)
                st.session_state.last_saved = time.time()
                st.success("Saved!")
        with c2:
            if st.button("📂 Load"):
                try:
                    d = json.load(open("session.json"))
                    st.session_state.book = d["book"]
                    st.session_state.chapter_order = d["chapters"]
                    st.session_state.characters = d.get("characters", [])
                    st.success("Loaded!")
                except Exception as e:
                    st.error(f"Load failed: {e}")
        with c3:
            if st.button("🆕 New"):
                for k in ("book","chapter_order","outline","cover","image_cache","characters","gen_progress"):
                    st.session_state[k] = {} if isinstance(st.session_state[k], dict) else []
                st.session_state.selected_genre = None
                st.session_state.selected_subgenres = []
                st.session_state.selected_tone = None
                st.session_state.last_saved = None

    # Logo
    logo_path = os.path.join(os.path.dirname(__file__), "logo.png")
    if os.path.exists(logo_path):
        st.markdown('<div class="logo-container">', unsafe_allow_html=True)
        st.image(logo_path, width=240)
        st.markdown('</div>', unsafe_allow_html=True)

    st.title("📖 NarrativaX – AI‑Powered Story Studio")

    # Model & presets
    st.subheader("🤖 AI Model")
    st.session_state.selected_model = st.radio(
        "", list(MODELS.keys()),
        index=list(MODELS.keys()).index(st.session_state.selected_model),
        horizontal=True, label_visibility="collapsed"
    )
    st.subheader("🖼️ Image Model")
    st.session_state.selected_image_model = st.radio(
        "", list(IMAGE_MODELS.keys()),
        index=list(IMAGE_MODELS.keys()).index(st.session_state.selected_image_model),
        horizontal=True, label_visibility="collapsed"
    )
    st.subheader("🔞 Content Intensity")
    preset = st.radio(
        "", list(PRESETS.keys()),
        index=list(PRESETS.keys()).index(st.session_state.content_preset),
        horizontal=True, label_visibility="collapsed"
    )
    st.session_state.content_preset = preset
    st.session_state.explicit_level = PRESETS[preset]

    # Genre & subgenres
    st.subheader("🎭 Genre & Sub‑Genres")
    genre = st.selectbox("Main Genre", ["-- choose --"] + list(GENRES.keys()))
    if genre in GENRES:
        st.session_state.selected_genre = genre
        subs = st.multiselect(
            "Pick sub‑genres", options=GENRES[genre],
            default=st.session_state.selected_subgenres
        )
        st.session_state.selected_subgenres = subs

    # Tone
    st.subheader("🎨 Tone")
    tone = st.selectbox("Narrative Tone", ["-- choose --"] + list(TONES.keys()))
    if tone in TONES:
        st.session_state.selected_tone = tone

    # Chapters & seed
    st.subheader("📑 Chapters & Seed")
    c1, c2 = st.columns([1, 3])
    with c1:
        chapters = st.slider("", 3, 30, 10, label_visibility="collapsed")
    with c2:
        prompt = st.text_input(
            "", placeholder="A dystopian romance between an AI and human rebel…",
            label_visibility="collapsed"
        )

    if st.button("🚀 Generate Book", use_container_width=True, type="primary"):
        if not (st.session_state.selected_genre and st.session_state.selected_tone and prompt.strip()):
            st.warning("Please choose genre, tone, and enter a seed.")
        else:
            st.session_state.gen_progress = {
                "prompt":    prompt.strip(),
                "genre":     st.session_state.selected_genre,
                "subgenres": st.session_state.selected_subgenres,
                "tone":      st.session_state.selected_tone,
                "chapters":  chapters,
                "model":     st.session_state.selected_model,
                "img_model": st.session_state.selected_image_model
            }
            generate_book_content()

    # Content tabs
    if st.session_state.chapter_order:
        tabs = st.tabs(["📖 Chapters","🎙️ Narration","🖼️ Artwork","📤 Export","👥 Characters"])

        # Chapters
        with tabs[0]:
            st.subheader("Chapter Manager")
            new_order = sort_items(st.session_state.chapter_order)
            if new_order:
                st.session_state.chapter_order = new_order

            for title in st.session_state.chapter_order:
                edit_flag = f"edit_{title}"
                if edit_flag not in st.session_state:
                    st.session_state[edit_flag] = False
                with st.expander(title):
                    if st.session_state[edit_flag]:
                        txt = st.text_area(
                            "", st.session_state.book[title], height=300,
                            label_visibility="collapsed"
                        )
                        if st.button("💾 Save", key=f"save_{title}"):
                            st.session_state.book[title] = txt
                            st.session_state[edit_flag] = False
                            st.success("Saved.")
                    else:
                        st.write(st.session_state.book[title])

                    c1, c2, c3 = st.columns(3)
                    with c1:
                        if st.button("➡️ Continue", key=f"cont_{title}"):
                            more = call_openrouter(
                                f"Continue chapter: {st.session_state.book[title]}",
                                st.session_state.selected_model
                            )
                            st.session_state.book[title] += "\n\n" + more
                    with c2:
                        if st.button("✏️ Edit", key=f"editbtn_{title}"):
                            st.session_state[edit_flag] = True
                    with c3:
                        if st.button("🗑️ Delete", key=f"del_{title}"):
                            st.session_state.book.pop(title, None)
                            st.session_state.chapter_order.remove(title)
                            st.success(f"Deleted {title}")

        # Narration
        with tabs[1]:
            st.subheader("Audio Narration")
            for title in st.session_state.chapter_order:
                with st.expander(f"🔊 {title}"):
                    txt = st.session_state.book[title]
                    try:
                        tts = gTTS(text=txt, lang="en")
                        with NamedTemporaryFile(suffix=".mp3") as fp:
                            tts.save(fp.name)
                            st.audio(fp.name)
                    except gTTSError:
                        st.error("TTS failed.")

        # Artwork
        with tabs[2]:
            st.subheader("Generated Artwork")
            if st.session_state.cover:
                st.image(st.session_state.cover, caption="Cover", use_container_width=True)
                if st.button("🌟 Regenerate Cover"):
                    cover = generate_image(
                        f"Cover for {st.session_state.gen_progress['prompt']}, "
                        f"{st.session_state.gen_progress['genre']} "
                        f"({', '.join(st.session_state.gen_progress['subgenres'])}), "
                        f"{st.session_state.gen_progress['tone']}",
                        st.session_state.selected_image_model, "cover"
                    )
                    st.session_state.cover = cover
            cols = st.columns(3)
            for i, (sec, img) in enumerate(st.session_state.image_cache.items()):
                with cols[i % 3]:
                    st.image(img, caption=sec.title(), use_container_width=True)

        # Export
        with tabs[3]:
            st.subheader("Export")
            if st.button("📦 Download ZIP", use_container_width=True):
                buf = create_export_zip()
                st.download_button(
                    "⬇️ Download ZIP",
                    data=buf.getvalue(),
                    file_name="narrativax_book.zip",
                    mime="application/zip"
                )

        # Characters
        with tabs[4]:
            st.subheader("Character Development")
            if st.button("➕ Add Character"):
                new_char = call_openrouter(
                    f"Create one character profile for seed: {st.session_state.gen_progress.get('prompt','')}",
                    st.session_state.selected_model
                ).strip()
                if new_char:
                    st.session_state.characters.append({"desc": new_char, "img": None})

            for idx, char in enumerate(st.session_state.characters):
                key_prefix = f"char_{idx}"
                with st.expander(f"🧑 Character #{idx+1}"):
                    c1, c2 = st.columns([3, 1])
                    with c1:
                        desc = st.text_area(
                            "", value=char["desc"], height=150,
                            key=f"{key_prefix}_desc", label_visibility="collapsed"
                        )
                        if desc != char["desc"]:
                            st.session_state.characters[idx]["desc"] = desc
                    with c2:
                        if st.button("🔄 Regenerate", key=f"{key_prefix}_regen"):
                            new_desc = call_openrouter(
                                f"Regenerate character profile: {char['desc']}",
                                st.session_state.selected_model
                            ).strip()
                            if new_desc:
                                st.session_state.characters[idx]["desc"] = new_desc
                        if st.button("🎨 Visualize", key=f"{key_prefix}_vis"):
                            img = generate_image(
                                char["desc"], st.session_state.selected_image_model, f"char_{idx}"
                            )
                            if img:
                                st.session_state.characters[idx]["img"] = img
                                st.image(img, use_container_width=True)
                        if st.button("🗑️ Delete", key=f"{key_prefix}_del"):
                            st.session_state.characters.pop(idx)
                            st.success("Deleted character")
                            break  # avoid index errors after deletion

if __name__ == "__main__":
    main_interface()
